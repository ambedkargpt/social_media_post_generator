from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from dotenv import load_dotenv
from pymongo import MongoClient


def main() -> None:
    load_dotenv()

    mongodb_uri = (os.getenv("MONGODB_URI") or "").strip()
    mongodb_database = (os.getenv("MONGODB_DATABASE") or "ambedkargpt").strip() or "ambedkargpt"
    if not mongodb_uri:
        raise ValueError("MONGODB_URI is not set. Please set it in your environment or .env file.")

    client = MongoClient(mongodb_uri)
    collection = client[mongodb_database]["questions"]

    questions = list(
        collection.find(
            {},
            {
                "_id": 0,
                "question_id": 1,
                "question_text": 1,
            },
        ).sort("created_at", 1)
    )

    if not questions:
        raise RuntimeError("No questions found in the database collection 'questions'.")

    doc = Document()
    doc.add_heading("Questions with Question IDs", level=1)

    for idx, item in enumerate(questions, start=1):
        qid = (item.get("question_id") or "").strip()
        qtext = (item.get("question_text") or "").strip()
        doc.add_paragraph(f"{idx}. {qid}")
        doc.add_paragraph(qtext if qtext else "(No question text)")
        doc.add_paragraph("")

    output_path = Path("outputs") / "questions_with_ids.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print(f"Exported {len(questions)} questions to: {output_path.resolve()}")


if __name__ == "__main__":
    main()
