from __future__ import annotations

import argparse
import subprocess
import time
import random
import re
import json
from pathlib import Path
from datetime import datetime, timezone

from tqdm import tqdm
from yt_dlp import YoutubeDL
from youtube_transcript_api import (
    YouTubeTranscriptApi,
    TranscriptsDisabled,
    NoTranscriptFound,
)

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


PROJECT_ROOT = Path(__file__).resolve().parent

from src.verification.transcript_verifier import TranscriptVerifier


INPUT_TXT = PROJECT_ROOT / "input.txt"
OUTPUT_BASE_DIR = PROJECT_ROOT / "outputs" / "transcripts"
DATA_DIR = PROJECT_ROOT / "data"
RAVISH_CHANNEL_SLUG = "ravishkumar.official"
RAVISH_DATA_TXT = DATA_DIR / "ravishkumar_all_transcripts.txt"
RAG_INDEX_PATH = PROJECT_ROOT / "outputs" / "faiss_index.bin"
RAG_CHUNKS_PATH = DATA_DIR / "argument_chunks.json"
RAG_VIDEO_CONTEXT_PATH = DATA_DIR / "video_context.json"
RAG_TITLE_EMB_PATH = DATA_DIR / "video_title_embeddings.json"
VIDEO_SUMMARIES_PATH = DATA_DIR / "video_summaries.json"

INPUT_TXT.parent.mkdir(parents=True, exist_ok=True)



MIN_DURATION_SEC = 0
MAX_DURATION_SEC = 1000 * 60 * 60

DOCX_FONT_NAME = "Mongolian Baiti"
TITLE_FONT_SIZE = 16
BODY_FONT_SIZE = 11
LINK_FONT_SIZE = 10

SLEEP_MIN = 30
SLEEP_MAX = 45
FETCH_DELAY = 2  # Delay in seconds between metadata fetches to avoid rate limiting



def sanitize_filename(text: str) -> str:
    cleaned = text.replace("\n", "_").replace("\r", "_")
    return re.sub(r'[\\/*?:"<>|]', "", cleaned)[:80]


def extract_channel_name(channel_url: str) -> str:
    # Handle playlist URLs
    if "playlist?" in channel_url or "list=" in channel_url:
        # Extract playlist ID
        match = re.search(r'list=([^&]+)', channel_url)
        if match:
            return f"playlist_{match.group(1)}"
        return "unknown_playlist"
    
    # Handle channel URLs
    match = re.search(r"@([^/]+)", channel_url)
    return match.group(1) if match else "unknown_channel"


def set_run_font(run, size, bold=False):
    run.font.name = DOCX_FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_NAME)


def ask_channel_url() -> str:
    url = input("Enter YouTube channel or playlist URL:\n> ").strip()
    if not url:
        raise ValueError("URL cannot be empty.")
    return url



def fetch_video_urls(channel_url: str, limit: int | None = None) -> list[str]:
    # Only convert channel URLs to videos tab, not playlists
    is_playlist = "playlist?" in channel_url or "list=" in channel_url
    
    if not is_playlist and "@" in channel_url:
        channel_id = channel_url.split("@")[1].split("/")[0]
        channel_url = f"https://www.youtube.com/@{channel_id}/videos"
    
    ydl_opts = {
        "flat_playlist": True,
        "quiet": False,
        "no_warnings": True,
        "socket_timeout": 30,
        "extract_flat": "in_playlist",
    }
    
    try:
        with YoutubeDL(ydl_opts) as ydl:
            print(" Fetching video list from {}...".format("playlist" if is_playlist else "channel"))
            info = ydl.extract_info(channel_url, download=False)
            entries = info.get("entries", [])
            
            # Filter and limit videos - stop after reaching limit
            urls = []
            for entry in entries:
                if entry.get("url") and "/shorts/" not in entry.get("url", ""):
                    urls.append(entry["url"])
                elif entry.get("webpage_url") and "/shorts/" not in entry.get("webpage_url", ""):
                    urls.append(entry["webpage_url"])
                
                # Stop if we've reached the limit
                if limit and limit > 0 and len(urls) >= limit:
                    urls = urls[:limit]
                    break
            
            source_label = "playlist" if is_playlist else "channel videos tab"
            print(f" Found {len(urls)} videos from {source_label} (shorts excluded)")
            return urls
    except Exception as e:
        print(f" Error fetching video URLs: {e}")
        print(" Tip: YouTube may have rate-limited your session. Try again in a few minutes.")
        return []


def _publish_meta_from_ytdlp(info: dict) -> dict:
    """
    YouTube metadata from yt-dlp: upload_date (YYYYMMDD), timestamp (Unix UTC) when available.
    """
    out: dict = {}
    ud = info.get("upload_date")
    if ud and isinstance(ud, str) and len(ud) == 8 and ud.isdigit():
        out["upload_date"] = ud
    ts = info.get("timestamp")
    if ts is not None:
        try:
            ts_i = int(ts)
            out["upload_timestamp"] = ts_i
            out["upload_datetime_utc"] = datetime.fromtimestamp(ts_i, tz=timezone.utc).isoformat()
        except (TypeError, ValueError, OSError):
            pass
    return out


def get_video_metadata(url: str) -> dict | None:
    ydl_opts = {
        "quiet": True,
        "no_warnings": True,
        "skip_download": True,
    }
    try:
        with YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            meta = {
                "id": info.get("id"),
                "title": info.get("title"),
                "duration": info.get("duration", 0),
                "url": url,
            }
            meta.update(_publish_meta_from_ytdlp(info))
            return meta
    except Exception as e:
        print(f" Metadata error: {e}")
        return None



def fetch_transcript_text(video_id: str) -> str | None:
    try:
        transcript = (
            YouTubeTranscriptApi()
            .fetch(video_id, languages=["hi", "en"])
            .to_raw_data()
        )
        return " ".join(t["text"] for t in transcript)
    except (TranscriptsDisabled, NoTranscriptFound):
        return None
    except Exception as e:
        print(f" Transcript error: {e}")
        return None



def create_docx(
    path: Path,
    title: str,
    url: str,
    transcript: str,
    *,
    upload_date: str | None = None,
    upload_datetime_utc: str | None = None,
) -> None:
    doc = Document()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_r = title_p.add_run(title)
    set_run_font(title_r, TITLE_FONT_SIZE, bold=True)

    link_p = doc.add_paragraph()
    link_r = link_p.add_run(f"Link: {url}")
    set_run_font(link_r, LINK_FONT_SIZE)

    if upload_datetime_utc:
        meta_p = doc.add_paragraph()
        meta_r = meta_p.add_run(f"Published (UTC): {upload_datetime_utc}")
        set_run_font(meta_r, LINK_FONT_SIZE)
    elif upload_date:
        meta_p = doc.add_paragraph()
        meta_r = meta_p.add_run(f"Upload date: {upload_date}")
        set_run_font(meta_r, LINK_FONT_SIZE)

    doc.add_paragraph("")

    body_p = doc.add_paragraph()
    body_r = body_p.add_run(transcript)
    set_run_font(body_r, BODY_FONT_SIZE)

    doc.save(path)


def create_txt_from_docx(docx_path: Path) -> None:
    """Read DOCX file and create a TXT file with the same content."""
    doc = Document(docx_path)
    txt_path = docx_path.with_suffix(".txt")
    
    with txt_path.open("w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n")


def consolidated_txt_name(channel_name: str) -> str:
    """
    Per-channel consolidated transcript filename.
    Keep Ravish channel name stable for existing pipeline.
    """
    if channel_name == RAVISH_CHANNEL_SLUG:
        return "ravishkumar_all_transcripts.txt"
    return f"{channel_name}_all_transcripts.txt"


def append_entries_to_consolidated(target_path: Path, entries: list[dict]) -> int:
    """
    Append only newly fetched transcript entries to the consolidated text file.
    Returns count of entries appended.
    """
    if not entries:
        return 0

    target_path.parent.mkdir(parents=True, exist_ok=True)
    appended = 0
    with target_path.open("a", encoding="utf-8") as f:
        for item in entries:
            title = (item.get("title") or "").strip()
            url = (item.get("url") or "").strip()
            transcript = (item.get("transcript") or "").strip()
            if not title or not transcript:
                continue
            f.write(f"===== {title} =====\n")
            f.write(f"{title}\n")
            f.write(f"Link: {url}\n")
            if item.get("upload_datetime_utc"):
                f.write(f"Published (UTC): {item['upload_datetime_utc']}\n")
            elif item.get("upload_date"):
                f.write(f"Upload date: {item['upload_date']}\n")
            f.write(f"{transcript}\n\n")
            appended += 1
    return appended


def rebuild_rag_artifacts_from_data_file(data_txt_path: Path) -> None:
    """
    Rebuild RAG artifacts after transcript data grows.
    This keeps chunks/index/title embeddings up to date automatically.
    """
    if not data_txt_path.exists():
        return

    from config import get_settings
    from pipeline.transcript_parser import parse_transcripts
    from pipeline.chunker import chunk_videos
    from pipeline.argument_scorer import score_argument_chunks
    from pipeline.embedder import ChunkEmbedder
    from pipeline.vector_store import build_index, save_vector_store
    from pipeline.title_embeddings import build_title_embeddings, save_title_embeddings

    settings = get_settings()
    raw_text = data_txt_path.read_text(encoding="utf-8")
    videos = parse_transcripts(raw_text)
    if not videos:
        return

    # Persist latest full video context
    RAG_VIDEO_CONTEXT_PATH.parent.mkdir(parents=True, exist_ok=True)
    RAG_VIDEO_CONTEXT_PATH.write_text(
        json.dumps(videos, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    chunks = chunk_videos(videos)
    chunks = score_argument_chunks(chunks)

    RAG_CHUNKS_PATH.parent.mkdir(parents=True, exist_ok=True)
    RAG_CHUNKS_PATH.write_text(
        json.dumps(chunks, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    embedder = ChunkEmbedder(
        api_key=settings.gemini_api_key,
        model_name=settings.embedding_model,
        batch_size=settings.embedding_batch_size,
    )
    embeddings, chunks = embedder.embed_chunks(
        chunks,
        cache_path=settings.embedding_chunk_cache_path,
        use_cache=settings.embedding_chunk_cache_enabled,
    )
    store = build_index(embeddings, chunks)
    save_vector_store(store, RAG_INDEX_PATH, RAG_CHUNKS_PATH)

    title_emb = build_title_embeddings(videos, embedder)
    save_title_embeddings(title_emb, RAG_TITLE_EMB_PATH)


def rebuild_semrag_artifacts_from_data_file(data_txt_path: Path) -> None:
    """
    Refresh SEMRAG chunks + KG from the latest transcript master file.
    Uses extraction cache, so only truly new/changed chunks are processed.
    """
    if not data_txt_path.exists():
        return

    from config import get_settings
    from pipeline.transcript_parser import parse_transcripts
    from semrag.build import build_semrag_graph, save_semrag_chunks
    from semrag.chunking import chunk_videos_for_semrag
    from semrag.semrag_config import load_semrag_config

    settings = get_settings()
    semrag_cfg = load_semrag_config(PROJECT_ROOT)
    semrag_chunks = []
    if settings.semrag_chunks_path.exists():
        try:
            semrag_chunks = json.loads(settings.semrag_chunks_path.read_text(encoding="utf-8"))
            if semrag_chunks:
                print(f" Using existing SEMRAG chunks from: {settings.semrag_chunks_path}")
        except Exception:
            semrag_chunks = []

    if not semrag_chunks:
        raw_text = data_txt_path.read_text(encoding="utf-8")
        videos = parse_transcripts(raw_text)
        if not videos:
            return
        semrag_chunks = chunk_videos_for_semrag(videos, embedder=None, cfg=semrag_cfg)
        save_semrag_chunks(settings.semrag_chunks_path, semrag_chunks)
    build_semrag_graph(
        chunks=semrag_chunks,
        settings=settings,
        graph_path=settings.semrag_graph_path,
        cache_path=settings.semrag_cache_path,
        force_rebuild=False,
    )


def build_semrag_graph_from_backups_only() -> None:
    """
    Rebuild semrag_graph.json strictly from extracted backup files without running extraction.
    """
    from config import get_settings
    from semrag.store import normalize_text, rebuild_indexes, save_semrag_graph

    settings = get_settings()
    semrag_dir = settings.semrag_graph_path.parent
    entities_path = semrag_dir / "semrag_entities_backup.json"
    relations_path = semrag_dir / "semrag_relations_backup.json"
    if not entities_path.exists() or not relations_path.exists():
        raise FileNotFoundError(
            "Missing backup files. Expected semrag_entities_backup.json and semrag_relations_backup.json."
        )

    entities_payload = json.loads(entities_path.read_text(encoding="utf-8"))
    relations_payload = json.loads(relations_path.read_text(encoding="utf-8"))
    entities = entities_payload.get("entities") if isinstance(entities_payload, dict) else []
    relations = relations_payload.get("relations") if isinstance(relations_payload, dict) else []
    if not isinstance(entities, list) or not isinstance(relations, list):
        raise ValueError("Backup files are malformed: entities/relations arrays are required.")

    graph = {
        "version": 1,
        "updated_at": entities_payload.get("updated_at") or relations_payload.get("updated_at"),
        "entities": entities,
        "relations": relations,
        "entity_name_to_id": {},
        "chunk_entities": {},
        "entity_to_chunks": {},
        "relation_to_chunks": {},
    }

    for ent in entities:
        if not isinstance(ent, dict):
            continue
        entity_id = str(ent.get("entity_id") or "").strip()
        name = str(ent.get("canonical_name") or "").strip()
        entity_type = str(ent.get("entity_type") or "").strip().lower()
        if entity_id and name and entity_type:
            graph["entity_name_to_id"][f"{normalize_text(name)}::{normalize_text(entity_type)}"] = entity_id

    chunk_entities: dict[str, set[str]] = {}
    for rel in relations:
        if not isinstance(rel, dict):
            continue
        cid = str(rel.get("evidence_chunk_id") or "").strip()
        if not cid:
            continue
        bucket = chunk_entities.setdefault(cid, set())
        head = str(rel.get("head_entity_id") or "").strip()
        tail = str(rel.get("tail_entity_id") or "").strip()
        if head:
            bucket.add(head)
        if tail:
            bucket.add(tail)
    graph["chunk_entities"] = {k: sorted(v) for k, v in chunk_entities.items()}
    rebuild_indexes(graph)
    save_semrag_graph(settings.semrag_graph_path, graph)


def summarize_fetched_entries(entries: list[dict], summary_path: Path, settings) -> tuple[int, list[dict]]:
    """
    Append DeepSeek summaries for fetch entries into data/video_summaries.json (cache by title+URL).
    Returns:
      - number of newly generated summaries
      - list of newly generated summary rows (title/link/summary_text + upload metadata)
    """
    if not entries:
        return 0, []

    from pipeline.video_summarizer import (
        deepseek_chat_client,
        get_or_create_video_summary,
        load_summary_cache,
        save_summary_cache,
        summary_cache_key,
    )

    client = deepseek_chat_client(settings)
    cache = load_summary_cache(summary_path)
    new_count = 0
    new_summary_rows: list[dict] = []
    print(f" Building video summaries for {len(entries)} new transcript(s)...")
    for item in tqdm(
        entries,
        desc="Summarizing videos",
        unit="video",
        dynamic_ncols=True,
    ):
        title = (item.get("title") or "").strip()
        url = (item.get("url") or "").strip()
        transcript = (item.get("transcript") or "").strip()
        if not title or not transcript:
            continue
        key = summary_cache_key(title, url)
        had_summary = bool(cache.get(key, {}).get("summary_text"))
        get_or_create_video_summary(
            client=client,
            model=settings.deepseek_summary_model,
            cache_entries=cache,
            video_title=title,
            video_link=url,
            full_text=transcript,
            target_words=190,
            prompts_dir=settings.prompts_dir,
        )
        if not had_summary:
            new_count += 1
            created = cache.get(key, {}) if isinstance(cache.get(key, {}), dict) else {}
            new_row = {
                "video_title": title,
                "video_link": url,
                "summary_text": str(created.get("summary_text", "")).strip(),
            }
            if item.get("upload_timestamp") is not None:
                new_row["upload_timestamp"] = item["upload_timestamp"]
            if item.get("upload_datetime_utc"):
                new_row["upload_datetime_utc"] = item["upload_datetime_utc"]
            if item.get("upload_date"):
                new_row["upload_date"] = item["upload_date"]
            if new_row["summary_text"]:
                new_summary_rows.append(new_row)
            save_summary_cache(summary_path, cache)
            if settings.summary_sleep_seconds > 0:
                time.sleep(settings.summary_sleep_seconds)
    print(
        " Summary step complete: "
        f"{new_count} new summary(ies), "
        f"{max(0, len(entries) - new_count)} already cached."
    )
    return new_count, new_summary_rows


def load_processed(path: Path) -> tuple[set[str], list[dict]]:
    if not path.exists():
        return set(), []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        records: list[dict] = []
        ids: set[str] = set()
        if isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    vid = item.get("id")
                    title = item.get("title", "")
                    url = item.get("url", "")
                else:
                    vid = item
                    title = ""
                    url = ""
                if vid:
                    ids.add(vid)
                    records.append({"id": vid, "title": title, "url": url})
        return ids, records
    except Exception:
        return set(), []


def save_processed(path: Path, records: list[dict]) -> None:
    # Deduplicate by video id and sort by title for readability
    dedup: dict[str, dict] = {}
    for rec in records:
        vid = rec.get("id")
        if not vid:
            continue
        entry = {
            "id": vid,
            "title": rec.get("title", ""),
            "url": rec.get("url", ""),
        }
        if rec.get("upload_date"):
            entry["upload_date"] = rec["upload_date"]
        if rec.get("upload_timestamp") is not None:
            entry["upload_timestamp"] = rec["upload_timestamp"]
        if rec.get("upload_datetime_utc"):
            entry["upload_datetime_utc"] = rec["upload_datetime_utc"]
        dedup[vid] = entry
    sorted_records = sorted(dedup.values(), key=lambda r: (r.get("title") or "", r["id"]))
    path.write_text(json.dumps(sorted_records, indent=2, ensure_ascii=False), encoding="utf-8")


def add_processed(processed_ids: set[str], processed_records: list[dict], meta: dict, url: str, path: Path) -> None:
    vid = meta.get("id") if meta else None
    if not vid:
        return
    title = meta.get("title", "")
    record = {"id": vid, "title": title, "url": url}
    if meta.get("upload_date"):
        record["upload_date"] = meta["upload_date"]
    if meta.get("upload_timestamp") is not None:
        record["upload_timestamp"] = meta["upload_timestamp"]
    if meta.get("upload_datetime_utc"):
        record["upload_datetime_utc"] = meta["upload_datetime_utc"]
    if vid not in processed_ids:
        processed_ids.add(vid)
        processed_records.append(record)
        save_processed(path, processed_records)
        return
    if not any(r.get("id") == vid for r in processed_records):
        processed_records.append(record)
        save_processed(path, processed_records)


def filter_already_downloaded_urls(video_urls: list[str], processed_records: list[dict], output_dir: Path) -> tuple[list[str], int]:
    """Filter out URLs that already have transcript files in output_dir."""
    if not video_urls or not processed_records:
        return video_urls, 0

    processed_urls_with_files: set[str] = set()
    for rec in processed_records:
        url = rec.get("url")
        title = rec.get("title", "")
        if not url or not title:
            continue

        base_name = sanitize_filename(title)
        docx_path = output_dir / f"{base_name}.docx"
        txt_path = output_dir / f"{base_name}.txt"
        if docx_path.exists() or txt_path.exists():
            processed_urls_with_files.add(url)

    if not processed_urls_with_files:
        return video_urls, 0

    filtered = [url for url in video_urls if url not in processed_urls_with_files]
    skipped = len(video_urls) - len(filtered)
    return filtered, skipped


def save_consolidated_verification(reports: list, output_dir: Path, channel_name: str) -> None:
    """Save consolidated verification report from in-memory reports."""
    from dataclasses import asdict
    
    # Convert VerificationReport objects to dicts
    reports_dicts = [asdict(r) if hasattr(r, '__dataclass_fields__') else r for r in reports]
    
    # Compute summary stats
    total = len(reports_dicts)
    passed = sum(1 for r in reports_dicts if r.get("summary", {}).get("status") == "PASS")
    review = sum(1 for r in reports_dicts if r.get("summary", {}).get("status") == "REVIEW")
    failed = sum(1 for r in reports_dicts if r.get("summary", {}).get("status") == "FAIL")
    avg_score = sum(r.get("summary", {}).get("score", 0) for r in reports_dicts) / total if total > 0 else 0
    
    aggregated = {
        "channel": channel_name,
        "generated_at": datetime.now().isoformat(),
        "summary": {
            "total_videos": total,
            "passed": passed,
            "review": review,
            "failed": failed,
            "average_score": round(avg_score, 3),
        },
        "videos": reports_dicts,
    }
    
    report_path = output_dir / "verification_report.json"
    report_path.write_text(
        json.dumps(aggregated, indent=2, ensure_ascii=False),
        encoding="utf-8"
    )
    print(f"\n Verification report saved: {report_path}")



def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Fetch YouTube channel/playlist transcripts into DOCX (Mongol font)."
    )
    parser.add_argument(
        "--channel",
        help="YouTube channel or playlist URL",
    )
    parser.add_argument(
        "--limit",
        type=int,
        help="Maximum number of videos to process",
    )
    parser.add_argument(
        "--start",
        type=int,
        default=1,
        help="Start from video number (1-indexed, default: 1)",
    )
    parser.add_argument(
        "--end",
        type=int,
        help="End at video number (1-indexed, inclusive)",
    )
    parser.add_argument(
        "--delay",
        type=float,
        default=FETCH_DELAY,
        help="Delay in seconds between metadata fetches to avoid rate limiting (default: %(default)s)",
    )
    parser.add_argument(
        "--rebuild-rag-only",
        action="store_true",
        help="Rebuild RAG (chunks, embeddings cache, FAISS, title embeddings) from "
        "data/ravishkumar_all_transcripts.txt and exit. Use after a failed/interrupted rebuild.",
    )
    parser.add_argument(
        "--rebuild-semrag-only",
        action="store_true",
        help="Rebuild SEMRAG chunks+KG from data/ravishkumar_all_transcripts.txt and exit.",
    )
    parser.add_argument(
        "--graph-only-from-backup",
        action="store_true",
        help="Build semrag_graph.json only from semrag_entities_backup.json and semrag_relations_backup.json.",
    )
    return parser.parse_args()


def legacy_main() -> int:
    args = parse_args()
    print("\n=== Fetch pipeline started ===")

    if args.rebuild_rag_only:
        print(" Rebuilding RAG from master transcript file (no fetch)…")
        rebuild_rag_artifacts_from_data_file(RAVISH_DATA_TXT)
        print(" RAG rebuild finished.")
        return 0

    if args.rebuild_semrag_only:
        print(" Rebuilding SEMRAG from master transcript file (no fetch)…")
        rebuild_semrag_artifacts_from_data_file(RAVISH_DATA_TXT)
        print(" SEMRAG rebuild finished.")
        return 0

    if args.graph_only_from_backup:
        print(" Rebuilding SEMRAG graph only from extracted backup files…")
        build_semrag_graph_from_backups_only()
        print(" SEMRAG graph-only rebuild finished.")
        return 0

    try:
        channel_url = args.channel or ask_channel_url()
    except Exception as e:
        print(f" Error: {e}")
        return 1

    channel_name = sanitize_filename(extract_channel_name(channel_url))
    output_dir = OUTPUT_BASE_DIR / channel_name
    output_dir.mkdir(parents=True, exist_ok=True)
    consolidated_path = output_dir / consolidated_txt_name(channel_name)
    processed_path = output_dir / "processed.json"
    processed_ids, processed_records = load_processed(processed_path)

    print(f"\n Channel    : {channel_name}")
    print(f" Output dir : {output_dir}")
    print(f" Delay/meta : {args.delay}s")
    print("================================\n")

    
    try:
        video_urls = fetch_video_urls(channel_url)
    except subprocess.CalledProcessError:
        print(" yt-dlp failed")
        return 1

    # Apply range filtering if explicitly requested
    if args.start != 1 or args.end is not None:
        start_idx = (args.start - 1) if args.start else 0  # Convert to 0-indexed
        end_idx = args.end if args.end else len(video_urls)
        video_urls = video_urls[start_idx:end_idx]
        print(f" Filtered to videos {args.start or 1}-{min(args.end or len(video_urls), len(video_urls))}")

    # Pre-filter videos that already exist in output directory
    video_urls, skipped_existing = filter_already_downloaded_urls(video_urls, processed_records, output_dir)
    if skipped_existing > 0:
        print(f" Skipping {skipped_existing} videos already present in output directory")

    # Apply limit after skipping existing files so reruns continue to new videos
    if args.limit and args.limit > 0:
        print(f" Applying limit: processing next {args.limit} pending videos")
        video_urls = video_urls[: args.limit]

    if not video_urls:
        print(" No new videos to process.")
        return 1

    
    with INPUT_TXT.open("w", encoding="utf-8") as f:
        for url in video_urls:
            f.write(url + "\n")

    print(f" Video links written to: {INPUT_TXT}")
    print(f" Total videos queued for fetch: {len(video_urls)}\n")

    # Initialize verifier once and collect all reports
    verifier = TranscriptVerifier()
    verification_reports = []
    newly_fetched_entries: list[dict] = []
    stat_meta_fail = 0
    stat_duration_skip = 0
    stat_already_processed = 0
    stat_transcript_missing = 0
    
    for url in tqdm(
        video_urls,
        desc="Fetching transcripts",
        unit="video",
        dynamic_ncols=True,
    ):
        time.sleep(args.delay)  # Add delay to avoid rate limiting
        meta = get_video_metadata(url)
        if not meta:
            stat_meta_fail += 1
            continue

        if not (MIN_DURATION_SEC <= meta["duration"] <= MAX_DURATION_SEC):
            stat_duration_skip += 1
            continue

        base_name = sanitize_filename(meta["title"])
        filename = base_name + ".docx"
        docx_path = output_dir / filename
        txt_path = docx_path.with_suffix(".txt")

        # Skip transcript fetch if already processed, but still verify if not verified
        transcript_exists = docx_path.exists() or txt_path.exists()
        
        if meta["id"] in processed_ids and transcript_exists:
            # Already processed; read existing transcript and verify it
            if txt_path.exists():
                with txt_path.open("r", encoding="utf-8") as f:
                    transcript = f.read()
                verification_report = verifier.verify(transcript, video_id=meta["id"])
                verification_reports.append(verification_report)
            
            add_processed(processed_ids, processed_records, meta, url, processed_path)
            stat_already_processed += 1
            continue

        transcript = fetch_transcript_text(meta["id"])
        if not transcript:
            stat_transcript_missing += 1
            continue

        create_docx(
            docx_path,
            meta["title"],
            url,
            transcript,
            upload_date=meta.get("upload_date"),
            upload_datetime_utc=meta.get("upload_datetime_utc"),
        )
        create_txt_from_docx(docx_path)

        # Verify transcript quality and add to reports list
        verification_report = verifier.verify(transcript, video_id=meta["id"])
        verification_reports.append(verification_report)

        # Track only truly new fetched transcripts for consolidated append.
        if meta["id"] not in processed_ids:
            entry = {
                "title": meta.get("title", ""),
                "url": url,
                "transcript": transcript,
            }
            if meta.get("upload_date"):
                entry["upload_date"] = meta["upload_date"]
            if meta.get("upload_timestamp") is not None:
                entry["upload_timestamp"] = meta["upload_timestamp"]
            if meta.get("upload_datetime_utc"):
                entry["upload_datetime_utc"] = meta["upload_datetime_utc"]
            newly_fetched_entries.append(entry)

        add_processed(processed_ids, processed_records, meta, url, processed_path)

        sleep_time = random.uniform(SLEEP_MIN, SLEEP_MAX)
        time.sleep(sleep_time)

    # Save single consolidated verification report
    print("\n Writing verification report...")
    save_consolidated_verification(verification_reports, output_dir, channel_name)

    # Append new entries to channel-level consolidated text file.
    added_count = append_entries_to_consolidated(consolidated_path, newly_fetched_entries)
    if added_count:
        print(f" Appended {added_count} new transcripts to: {consolidated_path}")
    else:
        print(" No newly fetched transcripts to append to consolidated file.")

    # Keep Ravish master dataset synced for RAG pipeline and rebuild artifacts.
    if channel_name == RAVISH_CHANNEL_SLUG and added_count > 0:
        mirror_added = append_entries_to_consolidated(RAVISH_DATA_TXT, newly_fetched_entries)
        print(f" Synced {mirror_added} entries to: {RAVISH_DATA_TXT}")
        print(" Rebuilding RAG artifacts from updated Ravish dataset...")
        rebuild_rag_artifacts_from_data_file(RAVISH_DATA_TXT)
        print(" RAG artifacts refreshed (chunks, FAISS index, title embeddings).")
        print(" Rebuilding SEMRAG artifacts from updated Ravish dataset...")
        rebuild_semrag_artifacts_from_data_file(RAVISH_DATA_TXT)
        print(" SEMRAG artifacts refreshed (semrag chunks + graph/cache).")
        from config import get_settings as _get_settings
        from pipeline.news_generator import update_generated_news_rolling

        fetch_settings = _get_settings()
        n_summaries, new_summary_rows = summarize_fetched_entries(
            newly_fetched_entries, VIDEO_SUMMARIES_PATH, fetch_settings
        )
        if n_summaries:
            print(f" Wrote {n_summaries} new video summary(ies) to: {VIDEO_SUMMARIES_PATH}")
            roll = update_generated_news_rolling(fetch_settings, new_summary_rows, show_progress=True)
            print(
                " Updated rolling news: "
                f"+{roll['generated']} generated, "
                f"{roll['active_total']} active, "
                f"{roll['evicted_to_legacy']} moved to legacy."
            )
        else:
            print(f" Video summaries up to date (no new entries or all already cached): {VIDEO_SUMMARIES_PATH}")
    elif channel_name != RAVISH_CHANNEL_SLUG:
        print(" Skipping RAG/news refresh because channel is not Ravish master channel.")

    print("\n Fetch run stats")
    print(f" - Total queued videos        : {len(video_urls)}")
    print(f" - New transcripts fetched    : {len(newly_fetched_entries)}")
    print(f" - Already processed+existing : {stat_already_processed}")
    print(f" - Metadata failures          : {stat_meta_fail}")
    print(f" - Duration filtered          : {stat_duration_skip}")
    print(f" - Transcript unavailable     : {stat_transcript_missing}")
    print(f" - Verified transcripts       : {len(verification_reports)}")

    print("\n All transcripts fetched successfully.")
    return 0


def main() -> int:
    import sys

    # Temporary compatibility shim: default orchestration is now run_pipeline.py.
    # Operators can still force the legacy path while rollout is ongoing.
    if "--legacy-fetch" in sys.argv:
        sys.argv = [arg for arg in sys.argv if arg != "--legacy-fetch"]
        return legacy_main()

    from run_pipeline import main as orchestrator_main

    passthrough_args = sys.argv[1:]
    return orchestrator_main(passthrough_args)


if __name__ == "__main__":
    raise SystemExit(main())
